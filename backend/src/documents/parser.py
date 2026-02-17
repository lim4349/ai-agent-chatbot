"""Document parser supporting multiple file formats."""

from __future__ import annotations

import csv
import json
import tempfile
from dataclasses import dataclass
from pathlib import Path


@dataclass
class DocumentSection:
    """A section of a parsed document."""

    content: str
    page: int | None = None
    heading: str | None = None
    section_type: str = "paragraph"  # "paragraph", "heading", "code", "table"


class DocumentParser:
    """Parser for multiple document formats."""

    def parse_from_bytes(self, content: bytes, file_type: str) -> list[DocumentSection]:
        """Parse document from bytes.

        Args:
            content: Raw file bytes
            file_type: Type of file (pdf, docx, txt, md, csv, json)

        Returns:
            List of DocumentSection objects

        """
        # For text-based files, decode and parse directly
        file_type = file_type.lower()

        if file_type in ("txt", "md", "csv", "json"):
            # Detect encoding
            text = self._decode_bytes(content)

            if file_type == "txt":
                return self._parse_text_content(text)
            elif file_type == "md":
                return self._parse_md_content(text)
            elif file_type == "csv":
                return self._parse_csv_content(text)
            else:  # json
                return self._parse_json_content(text)
        else:
            # For binary files (pdf, docx), write to temp file
            import os

            suffix = f".{file_type}"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            try:
                path = Path(tmp_path)
                if file_type == "pdf":
                    return self._parse_pdf(path)
                elif file_type == "docx":
                    return self._parse_docx(path)
                else:
                    raise ValueError(f"Unsupported file type: {file_type}")
            finally:
                os.unlink(tmp_path)

    def _decode_bytes(self, content: bytes) -> str:
        """Decode bytes to string trying common encodings."""
        encodings = ["utf-8", "cp949", "euc-kr", "latin-1"]

        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue

        return content.decode("utf-8", errors="replace")

    def _parse_text_content(self, text: str) -> list[DocumentSection]:
        """Parse plain text content."""
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        return [DocumentSection(content=para, section_type="paragraph") for para in paragraphs]

    def _parse_md_content(self, text: str) -> list[DocumentSection]:
        """Parse Markdown content."""
        sections = []
        lines = text.split("\n")
        current_heading = None
        i = 0

        while i < len(lines):
            line = lines[i]

            # Check for headings
            if line.startswith("#"):
                heading_text = line.lstrip("#").strip()
                current_heading = heading_text
                sections.append(
                    DocumentSection(
                        content=heading_text, heading=current_heading, section_type="heading"
                    )
                )
                i += 1

            # Check for code blocks
            elif line.strip().startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                code_content = "\n".join(code_lines)
                sections.append(
                    DocumentSection(
                        content=code_content, heading=current_heading, section_type="code"
                    )
                )
                i += 1

            # Regular paragraph
            elif line.strip():
                para_lines = [line]
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
                    para_lines.append(lines[i])
                    i += 1
                para_text = " ".join(para_lines).strip()
                if para_text:
                    sections.append(
                        DocumentSection(
                            content=para_text, heading=current_heading, section_type="paragraph"
                        )
                    )
            else:
                i += 1

        return sections

    def _parse_csv_content(self, text: str) -> list[DocumentSection]:
        """Parse CSV content from string."""
        import io

        sections = []
        f = io.StringIO(text)
        reader = csv.reader(f)
        headers = next(reader, None)

        if headers:
            sections.append(DocumentSection(content=" | ".join(headers), section_type="heading"))

        for row in reader:
            if row:
                if headers and len(headers) == len(row):
                    row_text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row, strict=True))
                else:
                    row_text = " | ".join(row)

                sections.append(DocumentSection(content=row_text, section_type="table"))

        return sections

    def _parse_json_content(self, text: str) -> list[DocumentSection]:
        """Parse JSON content from string."""
        data = json.loads(text)
        sections = []
        self._extract_json_text(data, sections)
        return sections

    def parse(self, file_path: str, file_type: str) -> list[DocumentSection]:
        """Parse a document and return a list of sections.

        Args:
            file_path: Path to the file to parse.
            file_type: Type of file (pdf, docx, txt, md, csv, json).

        Returns:
            List of DocumentSection objects.

        Raises:
            ValueError: If file_type is not supported.
            FileNotFoundError: If file does not exist.

        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower()

        if file_type == "pdf":
            return self._parse_pdf(path)
        elif file_type == "docx":
            return self._parse_docx(path)
        elif file_type == "txt":
            return self._parse_txt(path)
        elif file_type == "md":
            return self._parse_md(path)
        elif file_type == "csv":
            return self._parse_csv(path)
        elif file_type == "json":
            return self._parse_json(path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _parse_pdf(self, path: Path) -> list[DocumentSection]:
        """Parse PDF file using pdfplumber."""
        try:
            import pdfplumber
        except ImportError as err:
            raise ImportError(
                "pdfplumber is required for PDF parsing. Install with: pip install pdfplumber"
            ) from err

        sections = []
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    # Split by paragraphs but keep as single section per page for now
                    sections.append(
                        DocumentSection(
                            content=text.strip(), page=page_num, section_type="paragraph"
                        )
                    )
        return sections

    def _parse_docx(self, path: Path) -> list[DocumentSection]:
        """Parse DOCX file using python-docx."""
        try:
            from docx import Document as DocxDocument
        except ImportError as err:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install with: pip install python-docx"
            ) from err

        doc = DocxDocument(path)
        sections = []
        current_heading = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading
            if para.style.name.startswith("Heading"):
                current_heading = text
                sections.append(
                    DocumentSection(content=text, heading=current_heading, section_type="heading")
                )
            else:
                sections.append(
                    DocumentSection(content=text, heading=current_heading, section_type="paragraph")
                )

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))

            if table_text:
                sections.append(
                    DocumentSection(
                        content="\n".join(table_text), heading=current_heading, section_type="table"
                    )
                )

        return sections

    def _detect_encoding(self, path: Path) -> str:
        """Detect file encoding trying common encodings."""
        encodings = ["utf-8", "cp949", "euc-kr", "latin-1"]

        for encoding in encodings:
            try:
                with open(path, encoding=encoding) as f:
                    f.read()
                return encoding
            except UnicodeDecodeError:
                continue

        return "utf-8"  # Default fallback

    def _parse_txt(self, path: Path) -> list[DocumentSection]:
        """Parse plain text file with encoding detection."""
        encoding = self._detect_encoding(path)

        with open(path, encoding=encoding) as f:
            content = f.read()

        # Split into paragraphs
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

        return [DocumentSection(content=para, section_type="paragraph") for para in paragraphs]

    def _parse_md(self, path: Path) -> list[DocumentSection]:
        """Parse Markdown file preserving structure."""
        encoding = self._detect_encoding(path)

        with open(path, encoding=encoding) as f:
            content = f.read()

        sections = []
        lines = content.split("\n")
        current_heading = None
        i = 0

        while i < len(lines):
            line = lines[i]

            # Check for headings
            if line.startswith("#"):
                heading_text = line.lstrip("#").strip()
                current_heading = heading_text
                sections.append(
                    DocumentSection(
                        content=heading_text, heading=current_heading, section_type="heading"
                    )
                )
                i += 1

            # Check for code blocks
            elif line.strip().startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                code_content = "\n".join(code_lines)
                sections.append(
                    DocumentSection(
                        content=code_content, heading=current_heading, section_type="code"
                    )
                )
                i += 1

            # Regular paragraph
            elif line.strip():
                para_lines = [line]
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
                    para_lines.append(lines[i])
                    i += 1
                para_text = " ".join(para_lines).strip()
                if para_text:
                    sections.append(
                        DocumentSection(
                            content=para_text, heading=current_heading, section_type="paragraph"
                        )
                    )
            else:
                i += 1

        return sections

    def _parse_csv(self, path: Path) -> list[DocumentSection]:
        """Parse CSV file treating each row as a section."""
        encoding = self._detect_encoding(path)

        sections = []
        with open(path, encoding=encoding, newline="") as f:
            reader = csv.reader(f)
            headers = next(reader, None)

            if headers:
                # Add header as first section
                sections.append(
                    DocumentSection(content=" | ".join(headers), section_type="heading")
                )

            for row in reader:
                if row:
                    # Create a formatted row with headers if available
                    if headers and len(headers) == len(row):
                        row_text = " | ".join(
                            f"{h}: {v}" for h, v in zip(headers, row, strict=True)
                        )
                    else:
                        row_text = " | ".join(row)

                    sections.append(DocumentSection(content=row_text, section_type="table"))

        return sections

    def _parse_json(self, path: Path) -> list[DocumentSection]:
        """Parse JSON file extracting text fields."""
        encoding = self._detect_encoding(path)

        with open(path, encoding=encoding) as f:
            data = json.load(f)

        sections = []
        self._extract_json_text(data, sections)
        return sections

    def _extract_json_text(
        self, data: object, sections: list[DocumentSection], prefix: str = ""
    ) -> None:
        """Recursively extract text from JSON data."""
        if isinstance(data, dict):
            for key, value in data.items():
                new_prefix = f"{prefix}.{key}" if prefix else key
                if isinstance(value, str):
                    sections.append(
                        DocumentSection(
                            content=f"{new_prefix}: {value}",
                            heading=new_prefix,
                            section_type="paragraph",
                        )
                    )
                elif isinstance(value, dict | list):
                    self._extract_json_text(value, sections, new_prefix)
        elif isinstance(data, list):
            for i, item in enumerate(data):
                new_prefix = f"{prefix}[{i}]"
                if isinstance(item, str):
                    sections.append(
                        DocumentSection(
                            content=f"{new_prefix}: {item}",
                            heading=prefix,
                            section_type="paragraph",
                        )
                    )
                elif isinstance(item, dict | list):
                    self._extract_json_text(item, sections, new_prefix)
